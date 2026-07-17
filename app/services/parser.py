"""Parse PDF and Word documents, preserving page numbers as metadata."""

from __future__ import annotations

import logging
from pathlib import Path

from llama_index.core import Document

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}


def parse_document(file_path: Path) -> list[Document]:
    """Return LlamaIndex Documents with page_number metadata."""
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix}")

    if suffix == ".pdf":
        return _parse_pdf(file_path)
    return _parse_docx(file_path)


def _parse_pdf(file_path: Path) -> list[Document]:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    documents: list[Document] = []

    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()
        if not text:
            continue
        documents.append(
            Document(
                text=text,
                metadata={
                    "page_number": page_num,
                    "source": file_path.name,
                    "file_type": "pdf",
                },
            )
        )

    if not documents:
        raise ValueError(
            f"Could not extract text from {file_path.name}. "
            "Try a text-based (non-scanned) PDF."
        )

    return documents


def _parse_docx(file_path: Path) -> list[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    if not paragraphs:
        raise ValueError(
            f"Could not extract text from {file_path.name}. "
            "The Word document appears empty."
        )

    # Approximate page breaks: ~40 paragraphs per page for citation purposes
    page_size = 40
    documents: list[Document] = []
    for i in range(0, len(paragraphs), page_size):
        chunk_paras = paragraphs[i : i + page_size]
        page_num = (i // page_size) + 1
        documents.append(
            Document(
                text="\n\n".join(chunk_paras),
                metadata={
                    "page_number": page_num,
                    "source": file_path.name,
                    "file_type": "docx",
                },
            )
        )
    return documents
