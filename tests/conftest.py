"""Shared pytest fixtures for API integration tests."""

from __future__ import annotations

import os
from collections.abc import AsyncGenerator
from unittest.mock import MagicMock, patch

import pytest
import pytest_asyncio
from httpx import ASGITransport, AsyncClient
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

os.environ["DATABASE_URL"] = "sqlite+aiosqlite:///:memory:"
os.environ["LLM_PROVIDER"] = "ollama"
os.environ["EMBEDDING_PROVIDER"] = "local"

from app.config import get_settings  # noqa: E402
from app.db.models import Base  # noqa: E402
from app.db.session import get_db  # noqa: E402
from app.main import app  # noqa: E402
from app.services.cache import clear_cache  # noqa: E402


@pytest_asyncio.fixture
async def db_engine():
    engine = create_async_engine("sqlite+aiosqlite:///:memory:")
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    yield engine
    await engine.dispose()


@pytest_asyncio.fixture
async def db_session(db_engine) -> AsyncGenerator[AsyncSession, None]:
    factory = async_sessionmaker(db_engine, expire_on_commit=False)
    async with factory() as session:
        yield session


@pytest_asyncio.fixture
async def client(db_engine, tmp_path, monkeypatch) -> AsyncGenerator[AsyncClient, None]:
    monkeypatch.setenv("UPLOAD_DIR", str(tmp_path / "uploads"))
    monkeypatch.setenv("VECTOR_STORE_DIR", str(tmp_path / "vectors"))
    monkeypatch.setenv("DATABASE_URL", "sqlite+aiosqlite:///:memory:")
    get_settings.cache_clear()
    clear_cache()

    import app.db.session as session_module

    session_module._engine = None
    session_module._session_factory = None

    factory = async_sessionmaker(db_engine, expire_on_commit=False)

    async def override_get_db() -> AsyncGenerator[AsyncSession, None]:
        async with factory() as session:
            yield session

    app.dependency_overrides[get_db] = override_get_db

    mock_process = MagicMock()
    with patch("app.api.routes.documents.process_document", mock_process):
        transport = ASGITransport(app=app)
        async with AsyncClient(transport=transport, base_url="http://test") as ac:
            ac.mock_process = mock_process  # type: ignore[attr-defined]
            yield ac

    app.dependency_overrides.clear()
    session_module._engine = None
    session_module._session_factory = None
    get_settings.cache_clear()
    clear_cache()


@pytest.fixture
def sample_docx(tmp_path):
    from docx import Document as DocxDocument

    path = tmp_path / "manual.docx"
    doc = DocxDocument()
    doc.add_paragraph("Product Warranty")
    doc.add_paragraph("The warranty period is 2 years from the date of purchase.")
    doc.add_paragraph("Page 2 content about returns and refunds.")
    doc.save(str(path))
    return path
